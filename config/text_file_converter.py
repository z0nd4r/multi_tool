import os
import re
from tkinter import messagebox

# Для DOCX
from docx import Document
from docx.shared import Mm
# для PDF
import fitz
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

import chardet

font_normal = "./timesnewromanpsmt.ttf"
font_italic = "./timesnewromanps_italicmt.ttf"

pdfmetrics.registerFont(TTFont('TimesNewRoman', font_normal))
pdfmetrics.registerFont(TTFont('TimesNewRoman-Italic', font_italic))

def convert_text(combo, file_path, output_path):
    text = None
    global input_file
    input_file = file_path.get()
    output_dir = output_path.get()

    # with open(input_file, 'rb') as f:
    #     text = f.read()
    #     result = chardet.detect(text)
    #     encoding = result["encoding"]
    #     print(f'Файл {input_file} имеет кодировку: {encoding}')

    choice = combo.get() # получаем выводимый формат

    if not input_file or not output_dir:
        messagebox.showerror("Ошибка", "Пожалуйста, выберите текстовый файл и папку для сохранения")
        return
    if not choice:
        messagebox.showerror("Ошибка", "Пожалуйста, выберите выводимый формат")
        return

    filename, ext = os.path.splitext(os.path.basename(input_file)) # забираем у входного файла его имя и формат
    output_file = os.path.join(output_dir, f'{filename}.{choice.lower()}') # новый файл с тем же именем, но с новым ф.

    if ext.upper() == '.'+choice:
        messagebox.showinfo('Внимание', 'Входной и выходной форматы совпадают')
        return

    # определяем входной формат
    if ext.upper() == '.DOCX':
        if choice == 'PDF':
            text = None
        else:
            text = get_text_from_docx(input_file)
    elif ext.upper() == '.PDF':
        text = get_text_from_pdf(input_file)
    elif ext.upper() == '.TXT':
        if choice == 'PDF':
            text = None
        else:
            text = get_text_from_txt(input_file)

    # если текст не найден, останавливаем программу
    if text is None and choice != 'PDF':
        return

    # проверка, есть ли файл с таким же именем в директории
    if os.path.exists(output_file):
        examination_and_save(text, output_file, output_dir, filename, choice, ext)
        return
    else:
        save_text_to_file(text, output_file, choice, ext)


def save_text_to_file(text, output_file, choice, ext):
    """
    Сохраняет текст в файл
    """
    try:
        if choice == 'TXT':
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)
        elif choice == 'PDF':
            if ext.upper() == '.DOCX':
                success, e = docx_to_pdf(input_file, output_file)
            else:
                print(os.path.exists(input_file))
                success, e = txt_to_pdf(input_file, output_file)
            if not success:
                print(e)
                messagebox.showerror('Ошибка', f'Ошибка при записи PDF файла:\n{e}')
                return
        elif choice == 'DOCX':
            success, e = write_text_to_docx(text, output_file)
            if not success:
                messagebox.showerror('Ошибка', f'Ошибка при записи DOCX файла:\nзакройте существующий DOCX файл и попробуйте снова')
                return

        messagebox.showinfo('Успех', f'Текстовый файл сохранен как:\n{output_file}')
    except Exception as e:
        print(e)
        messagebox.showerror('Ошибка', f'Ошибка при записи в файл:\n{e}')
        return


def get_text_from_docx(docx_file):
    """
    Извлекает текст из DOCX файла
    """
    try:
        document = Document(docx_file)
        full_text = ''

        for paragraph in document.paragraphs:

            # обрабатываем каждый параграф
            text = paragraph.text

            # Сохраняем выравнивание через табуляцию
            if paragraph.paragraph_format.first_line_indent:
                indent = paragraph.paragraph_format.first_line_indent
                if indent > 0:
                    text = '\t' + text

            # Обрабатываем стили с отступами
            if hasattr(paragraph, 'style') and paragraph.style:
                if 'Indent' in paragraph.style.name:
                    text = '\t' + text

            full_text += (text + '\n')

        return full_text

    except Exception as e:
        messagebox.showerror('Ошибка', f'Ошибка при чтении DOCX-файла:\n{e}')
        return False, e


def get_text_from_pdf(pdf_file):
    """
    Извлекает текст из PDF файла
    """
    try:
        pdf_document = fitz.open(pdf_file)
        full_text = ''
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            blocks = page.get_text('blocks')
            for block in blocks:
                text = block[4]  # текст находится в 5-м элементе кортежа
                full_text += text + "\n"  # добавляем перенос строки между блоками

        return full_text
    except Exception as e:
        messagebox.showerror('Ошибка', f'Ошибка при чтении PDF файла:\n{e}')
        return


def get_text_from_txt(txt_file):
    """
    Извлекает текст из TXT файла
    """
    try:
        with open(txt_file, 'r', encoding='utf-8') as f:
            full_text = f.read()
        return full_text
    except Exception as e:
        messagebox.showerror('Ошибка', f'Ошибка при чтении TXT файла:\n{e}')


def write_text_to_docx(text, output_file):
    """
    Создает DOCX файл из текста
    """
    try:
        document = Document()
        section = document.sections[0]

        section.top_margin = Mm(15)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(15)

        font_size = Pt(16)
        font_style = 'TimesNewRoman'

        for paragraph_text in text.splitlines():
            if '\t' in paragraph_text:
                # Для строк с табуляцией создаем отдельный параграф
                p = document.add_paragraph()
                for i, part in enumerate(paragraph_text.split('\t')):
                    run = p.add_run(part)
                    run.font.name = font_style
                    run.font.size = font_size
                    if i < len(paragraph_text.split('\t')) - 1:
                        # Добавляем табуляцию
                        run = p.add_run()
                        run.add_tab()
                        run.font.name = font_style
                        run.font.size = font_size
            else:
                p = document.add_paragraph()
                run = p.add_run(paragraph_text)
                run.font.name = font_style
                run.font.size = font_size

        document.save(output_file)
        return True, None
    except Exception as e:
        return False, e


def txt_to_pdf(txt_file, output_file):
    """
    Конвертирует TXT в PDF
    """
    try:
        # Создаём PDF-файл
        doc_pdf = SimpleDocTemplate(output_file, pagesize=A4,
                                    leftMargin=20 * mm, rightMargin=20 * mm,
                                    topMargin=20 * mm, bottomMargin=20 * mm)

        # Читаем текст
        with open(txt_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        style = ParagraphStyle(
            name='CustomStyle',
            fontName='TimesNewRoman',
            fontSize=12,
            leading=14,
            leftIndent=0,
            firstLineIndent=0,
            spaceBefore=0,
            spaceAfter=0,
            alignment=TA_LEFT
        )

        # Формируем контент
        content = []
        for line in lines:
            if line.strip():  # Пропускаем пустые строки
                content.append(Paragraph(line.strip(), style))
                content.append(Spacer(1, 5))  # Небольшой отступ между абзацами

        doc_pdf.build(content)
        return True, None
    except Exception as e:
        print(e)
        return False, e


def docx_to_pdf(docx_path, output_file):
    """
    Конвертирует DOCX в PDF с сохранением позиционирования и форматирования
    """
    try:
        # print("Текст:")
        # print(text)
        # print("\nСимволы разделителей абзацев:")
        # for i, char in enumerate(text):
        #     if ord(char) < 32:  # проверяем управляющие символы
        #         print(f"Символ {i}: {ord(char)}")
        # # print(text.split("\n\n"))

        # for fontname in pdfmetrics.standardFonts:
        #     print(fontname)

        # Загрузка DOCX документа
        doc = Document(docx_path)

        # Создание документа PDF
        doc_pdf = SimpleDocTemplate(output_file, pagesize=A4,
                                    leftMargin=20 * mm, rightMargin=20 * mm,
                                    topMargin=20 * mm, bottomMargin=20 * mm)

        story = []

        # Обработка всех элементов DOCX
        for element in doc.element.body:
            process_element(element, story)

        # Сборка PDF
        doc_pdf.build(story)
        return True, None
    except Exception as e:
        return False, e

def process_element(element, story):
    """
    Обрабатывает различные элементы DOCX
    """
    # Определяем тип элемента и обрабатываем соответствующим образом
    if element.tag.endswith('p'):  # Параграф
        process_paragraph(element, story)
    elif element.tag.endswith('tbl'):  # Таблица
        process_table(element, story)
    elif element.tag.endswith('sectPr'):  # Раздел (разрыв страницы)
        story.append(PageBreak())


def process_paragraph(element, story):
    """
    Обрабатывает параграф с сохранением форматирования
    """

    # Получаем свойства форматирования
    p_pr = element.find('.//w:pPr', namespaces=element.nsmap)

    # Определяем ВСЕ отступы и выравнивание
    left_indent = 0        # Отступ всего абзаца слева
    first_line_indent = 0  # Красная строка (отступ первой строки)
    space_before = 0       # Отступ перед абзацем
    space_after = 0        # Отступ после абзаца
    alignment = TA_LEFT    # Выравнивание

    if p_pr is not None:
        # Определяем выравнивание
        jc = p_pr.find('.//w:jc', namespaces=element.nsmap)
        if jc is not None:
            align_val = jc.get(qn('w:val'))
            if align_val == 'center':
                alignment = TA_CENTER
            elif align_val == 'right':
                alignment = TA_RIGHT
            elif align_val == 'both':
                alignment = TA_JUSTIFY

        # Определяем отступы абзаца
        ind = p_pr.find('.//w:ind', namespaces=element.nsmap)
        if ind is not None:
            if ind.get('left'):
                left_indent = float(ind.get('left')) / 20  # twips -> points
            if ind.get('firstLine'):
                first_line_indent = float(ind.get('firstLine')) / 20
            elif ind.get('hanging'):
                first_line_indent = -float(ind.get('hanging')) / 20

        # Определяем интервалы
        spacing = p_pr.find('.//w:spacing', namespaces=element.nsmap)
        if spacing is not None:
            if spacing.get('before'):
                space_before = float(spacing.get('before')) / 20
            if spacing.get('after'):
                space_after = float(spacing.get('after')) / 20

    # Получаем текст параграфа
    text = ''
    for run in element.findall('.//w:r', namespaces=element.nsmap):
        t = run.find('.//w:t', namespaces=element.nsmap)
        if t is not None:
            text += t.text or ''

    # Создаем стиль с точными отступами
    style = ParagraphStyle(
        name='CustomStyle',
        fontName='TimesNewRoman',
        fontSize=12,
        leading=14,
        leftIndent=left_indent,          # Отступ всего абзаца
        firstLineIndent=first_line_indent, # Красная строка
        spaceBefore=space_before,        # Отступ перед абзацем (ВЕРНУЛ!)
        spaceAfter=space_after,          # Отступ после абзаца (ВЕРНУЛ!)
        alignment=alignment
    )

    if text.strip():
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 2 * mm))


def process_table(element, story):
    """
    Обрабатывает таблицы с сохранением структуры
    """
    table_data = []

    for row in element.findall('.//w:tr', namespaces=element.nsmap):
        row_data = []
        for cell in row.findall('.//w:tc', namespaces=element.nsmap):
            cell_text = ''
            for p in cell.findall('.//w:p', namespaces=element.nsmap):
                for run in p.findall('.//w:r', namespaces=element.nsmap):
                    t = run.find('.//w:t', namespaces=element.nsmap)
                    if t is not None:
                        cell_text += t.text or ''
            row_data.append(cell_text)
        table_data.append(row_data)

    # Создаем таблицу
    if table_data:
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (-1, -1), 'TimesNewRoman'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('LEADING', (0, 0), (-1, -1), 12),
        ]))
        story.append(table)
        story.append(Spacer(1, 5 * mm))

def examination_and_save(text, output_file, output_dir, filename, choice, ext):
    """
    Выводит диалоговое окно, если файл существует в выбранной директории и сохранение файла
    """
    response = messagebox.askyesnocancel(
        title='Конфликт файла',
        message=f"Файл {os.path.basename(output_file)} уже существует!\n"
                "Хотите перезаписать его?",
        detail="Нажмите\n"
               "• 'Да' — перезаписать\n"
               "• 'Нет' — сохранить с новым именем\n"
               "• 'Отмена' — не сохранять",
        icon=messagebox.WARNING
    )

    if response is True:
        save_text_to_file(text, output_file, choice, ext)
    elif response is False:
        counter = 1
        while True:
            new_filename = f"{filename}_{counter}.{choice.lower()}"
            new_output_file = os.path.join(output_dir, new_filename)
            if not os.path.exists(new_output_file):
                save_text_to_file(text, new_output_file, choice, ext)
                break
            counter += 1
    else:
        messagebox.showinfo("Внимание", "Сохранение отменено")
        return





